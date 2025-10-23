import io
import os
import traceback
from fastapi import APIRouter, BackgroundTasks, HTTPException, UploadFile, File, Depends
from openai import BaseModel
from sqlalchemy.orm import Session
from typing import Dict
import time
from datetime import datetime
import uuid
import logging

# Schemas
from app.schemas.rag_schemas import (
    RAGQueryRequest,
    RAGQueryResponse,
    HealthCheckResponse,
    DocumentUploadResponse
)

# Database
from app.core.config import get_db
from app.models.rag_model import (
    Document,
    DocumentChunk,
    GraphEntity,
    GraphRelationship,
    Query as QueryModel,
    Session as SessionModel
)

# Orchestrator service
from app.core.dependencies import get_rag_service
from app.core.enums import RAGStrategy
from app.services.chunking import chunk_text
# from app.services.vectorstore import get_vector_store
from app.core.dependencies import get_rag_service

from app.models.rag_model import Query, Document, Session
logger = logging.getLogger(__name__)

# ✅ Router prefix
router = APIRouter(prefix="/api/rag", tags=["RAG System"])


# ============================================================
# 🧠 RAG Query Endpoint
# ============================================================
@router.post("/query", response_model=RAGQueryResponse)
async def query_rag(request: RAGQueryRequest, db: Session = Depends(get_db)):
    """Main RAG query endpoint with strategy support."""
    try:
        logger.info(f"[QUERY] Received query: {request.query[:50]}...")
        start_time = time.time()

        rag_service = get_rag_service()

        # ✅ Map strategy string to enum
        strategy_map = {
            "simple": RAGStrategy.SIMPLE,
            "agentic": RAGStrategy.AGENTIC,
            "auto": RAGStrategy.AUTO
        }

        if request.strategy not in strategy_map:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid strategy: {request.strategy}. Use: simple, agentic, or auto"
            )

        strategy = strategy_map[request.strategy]

        # ✅ Validate document exists if document_id is provided
        document_id = request.document_id
        document = None
        if document_id:
            document = db.query(Document).filter(Document.id == document_id).first()
            if not document:
                raise HTTPException(
                    status_code=404,
                    detail=f"Document with ID {document_id} not found"
                )
            logger.info(f"[QUERY] Querying document: {document.filename}")

        # ✅ Execute orchestrator with document_id
        result = await rag_service.execute_query(
            query=request.query,
            top_k=request.top_k,
            session_id=request.session_id,
            document_id=document_id,  # Pass document_id to RAG service
            strategy=strategy
        )

        processing_time = time.time() - start_time

        # ✅ Save query in DB with document_id
        db_query = Query(
            id=str(uuid.uuid4()),
            query_text=request.query,
            answer=result["answer"],
            strategy_used=request.strategy,
            processing_time=processing_time,
            confidence_score=result.get("confidence", 0.85),
            retrieved_chunks_count=len(result.get("retrieved_chunks", [])),
            session_id=request.session_id,
            document_id=document_id,  # Save document_id
            metadata={
                "top_k": request.top_k,
                "chunk_count": len(result.get("retrieved_chunks", [])),
                "document_id": document_id,
                "document_filename": document.filename if document else None
            }
        )
        db.add(db_query)
        db.commit()
        db.refresh(db_query)

        return RAGQueryResponse(
            query=request.query,
            answer=result["answer"],
            strategy_used=result["strategy_used"].value,
            processing_time=processing_time,
            retrieved_chunks=result.get("retrieved_chunks", []),
            confidence_score=result.get("confidence", 0.85)
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[ERROR] Query processing failed: {str(e)}", exc_info=True)
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Query processing failed: {str(e)}")


# ============================================================
# 📤 Document Upload Endpoint
# ============================================================
from PyPDF2 import PdfReader

try:
    import docx  # python-docx
except Exception:
    docx = None


def extract_text_from_pdf_bytes(pdf_bytes):
    """Extract text from PDF bytes with enhanced error handling"""
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        text = ""
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
                text += page_text
                logger.debug(f"[DEBUG] Extracted {len(page_text)} chars from page {page_num + 1}")
            except Exception as e:
                logger.warning(f"[WARN] Failed to extract text from page {page_num + 1}: {str(e)}")
                continue
        return text
    except Exception as e:
        logger.error(f"[ERROR] PDF extraction failed: {str(e)}")
        raise RuntimeError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes with enhanced error handling"""
    if docx is None:
        raise RuntimeError("python-docx is not installed. Install with: pip install python-docx")
    
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        text = "\n".join(paragraphs)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        text += "\n" + cell.text
        
        return text.strip()
    except Exception as e:
        logger.error(f"[ERROR] DOCX extraction failed: {str(e)}")
        raise RuntimeError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_txt_bytes(file_bytes: bytes) -> str:
    """Extract text from TXT bytes with enhanced encoding detection"""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'utf-16']
    
    for encoding in encodings:
        try:
            text = file_bytes.decode(encoding)
            # Validate that we got meaningful text
            if text.strip() and any(char.isalnum() for char in text):
                logger.debug(f"[DEBUG] Successfully decoded with {encoding}")
                return text
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.warning(f"[WARN] Encoding {encoding} failed: {str(e)}")
            continue
    
    # Fallback: decode with utf-8 ignoring errors
    logger.warning("[WARN] All encodings failed, using utf-8 with ignore errors")
    return file_bytes.decode("utf-8", errors="ignore")


# ✅ Update the response model to include processing_time as string
class DocumentUploadResponse(BaseModel):
    document_id: str
    filename: str
    status: str
    chunks_created: int
    message: str
    processing_time: str


@router.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """Upload and process document for RAG (PDF, TXT, DOCX) with enhanced monitoring and error handling."""
    start_time = time.time()
    processing_metrics = {
        'extraction_time': 0,
        'chunking_time': 0,
        'db_operations_time': 0,
        'vector_store_time': 0,
        'total_chars_processed': 0
    }
    
    try:
        logger.info(f"[UPLOAD] Received file upload: {file.filename} (Size: {file.size or 'unknown'})")

        # Validate file type
        allowed_types = [
            "application/pdf",
            "text/plain",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]
        
        if file.content_type not in allowed_types:
            # Also check by file extension as fallback
            file_extension = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
            if file_extension not in ['pdf', 'txt', 'docx']:
                raise HTTPException(
                    status_code=400,
                    detail=f"File type not supported. Allowed: PDF, TXT, DOCX. Got: {file.content_type}"
                )
            else:
                logger.info(f"[INFO] Using file extension {file_extension} for content type detection")

        # Read bytes
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Uploaded file is empty")

        file_size_mb = len(content) / (1024 * 1024)
        logger.info(f"[FILE] File size: {len(content)} bytes ({file_size_mb:.2f} MB)")

        # Extract text based on content type with timing
        extraction_start = time.time()
        try:
            if file.content_type == "application/pdf":
                text = extract_text_from_pdf_bytes(content)
            elif file.content_type == "text/plain":
                text = extract_text_from_txt_bytes(content)
            elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = extract_text_from_docx_bytes(content)
            else:
                # Fallback: try to detect by extension
                if file.filename.lower().endswith('.pdf'):
                    text = extract_text_from_pdf_bytes(content)
                elif file.filename.lower().endswith('.docx'):
                    text = extract_text_from_docx_bytes(content)
                else:
                    text = extract_text_from_txt_bytes(content)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Text extraction failed: {str(e)}")
        
        processing_metrics['extraction_time'] = time.time() - extraction_start

        # Enhanced text validation
        if not text or text.strip() == "":
            logger.warning(f"[WARN] No extractable text found in {file.filename}, content type: {file.content_type}")
            raise HTTPException(status_code=400, detail="No extractable text found in uploaded document")

        # Check if text is meaningful (not just whitespace/formatting)
        clean_text = text.strip()
        if len(clean_text) < 10:
            logger.warning(f"[WARN] Very little text extracted from {file.filename}: {len(clean_text)} chars")
            raise HTTPException(status_code=400, detail="Extracted text appears to be too short or invalid")

        processing_metrics['total_chars_processed'] = len(clean_text)
        logger.info(f"[TEXT] Extracted {len(clean_text)} characters from document")

        # Chunk the extracted text with timing
        chunking_start = time.time()
        chunks = chunk_text(clean_text)
        processing_metrics['chunking_time'] = time.time() - chunking_start

        # Validate chunks
        if not isinstance(chunks, list):
            raise HTTPException(status_code=500, detail="chunk_text did not return a list of chunks")

        if not chunks:
            raise HTTPException(status_code=400, detail="No chunks generated from document text")

        if len(chunks) == 0:
            raise HTTPException(status_code=400, detail="Empty chunks list generated")

        logger.info(f"[CHUNKS] Generated {len(chunks)} chunks for file {file.filename}")

        # Calculate chunk statistics
        chunk_sizes = []
        for chunk in chunks:
            chunk_text_content = chunk.get("content") if isinstance(chunk, dict) else chunk
            if chunk_text_content:
                chunk_sizes.append(len(chunk_text_content))

        if chunk_sizes:
            logger.info(f"[STATS] Chunk sizes: min={min(chunk_sizes)}, max={max(chunk_sizes)}, avg={sum(chunk_sizes)/len(chunk_sizes):.0f}")

        # ✅ Create document ID and DB record within transaction
        doc_id = str(uuid.uuid4())
        db_operations_start = time.time()
        
        try:
            db_doc = Document(
                id=doc_id,
                filename=file.filename,
                content_type=file.content_type,
                size=len(content),
                status="processing",
                chunks_count=0,
                uploaded_at=datetime.utcnow(),
                meta_data={
                    "original_size": len(content),
                    "processing_method": "local_extraction",
                    "extracted_chars": len(clean_text),
                    "chunks_generated": len(chunks)
                }
            )
            db.add(db_doc)
            db.flush()

            # Add chunks to vector store and database
            vector_store = get_vector_store()
            vector_store_start = time.time()
            
            chunks_created = 0
            failed_chunks = 0
            chunk_details = []

            for idx, chunk in enumerate(chunks):
                try:
                    chunk_text_content = chunk.get("content") if isinstance(chunk, dict) else chunk
                    if not chunk_text_content or not chunk_text_content.strip():
                        logger.warning(f"[WARN] Skipping empty chunk at index {idx}")
                        failed_chunks += 1
                        continue

                    chunk_id = str(uuid.uuid4())

                    # Create database chunk record
                    db_chunk = DocumentChunk(
                        id=chunk_id,
                        document_id=doc_id,
                        content=chunk_text_content,
                        chunk_index=idx,
                        meta_data={
                            "source": file.filename,
                            "content_type": file.content_type,
                            "chunk_size": len(chunk_text_content),
                            "chunk_chars": len(chunk_text_content)
                        },
                        created_at=datetime.utcnow()
                    )
                    db.add(db_chunk)

                    # Add to vector store
                    vector_store.add_document(
                        text=chunk_text_content,
                        metadata={
                            "chunk_id": chunk_id,
                            "chunk_index": idx,
                            "document_id": doc_id,
                            "source": file.filename,
                            "content_type": file.content_type,
                            "chunk_size": len(chunk_text_content)
                        }
                    )
                    chunks_created += 1
                    chunk_details.append({
                        "index": idx,
                        "size": len(chunk_text_content),
                        "status": "success"
                    })

                except Exception as e:
                    logger.error(f"[ERROR] Failed to process chunk {idx}: {str(e)}")
                    failed_chunks += 1
                    chunk_details.append({
                        "index": idx,
                        "size": len(chunk_text_content) if chunk_text_content else 0,
                        "status": "failed",
                        "error": str(e)
                    })
                    continue

            processing_metrics['vector_store_time'] = time.time() - vector_store_start

            # Update document record
            db_doc.status = "completed" if chunks_created > 0 else "failed"
            db_doc.chunks_count = chunks_created
            db_doc.processed_at = datetime.utcnow()
            db_doc.meta_data.update({
                "chunks_created": chunks_created,
                "chunks_failed": failed_chunks,
                "processing_metrics": processing_metrics,
                "chunk_details": chunk_details
            })

            processing_metrics['db_operations_time'] = time.time() - db_operations_start

            # Commit transaction
            db.commit()
            db.refresh(db_doc)

        except Exception as e:
            db.rollback()
            logger.error(f"[ERROR] Database transaction failed: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Database operation failed: {str(e)}")

        # Log final statistics
        total_processing_time = time.time() - start_time
        logger.info(f"[SUCCESS] Upload complete: {chunks_created} chunks created, {failed_chunks} failed")
        logger.info(f"[TIME] Processing times - Total: {total_processing_time:.2f}s, "
                   f"Extraction: {processing_metrics['extraction_time']:.2f}s, "
                   f"Chunking: {processing_metrics['chunking_time']:.2f}s, "
                   f"DB: {processing_metrics['db_operations_time']:.2f}s, "
                   f"Vector: {processing_metrics['vector_store_time']:.2f}s")

        # Optional: get vector count
        try:
            vector_count = vector_store.get_count()
            logger.info(f"[VECTOR] Vector store now contains {vector_count} total vectors")
        except Exception as e:
            logger.warning(f"[WARN] Could not retrieve vector count: {str(e)}")

        # ✅ Format processing time as "4.4 seconds"
        formatted_time = f"{total_processing_time:.1f} seconds"

        # ✅ Return the response with processing_time as formatted string
        return DocumentUploadResponse(
            document_id=db_doc.id,
            filename=file.filename,
            status="success",
            chunks_created=chunks_created,
            message=f"Document '{file.filename}' processed successfully with {chunks_created} chunks",
            processing_time=formatted_time
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[ERROR] Document upload failed: {str(e)}", exc_info=True)
        try:
            db.rollback()
        except Exception:
            pass  # Already rolled back or no transaction
        raise HTTPException(status_code=500, detail=f"Document upload failed: {str(e)}")
    
# ============================================================
# 🧹 Clear All Documents Endpoint
# ============================================================
@router.delete("/documents/clear")
async def clear_all_documents(db: Session = Depends(get_db)):
    """
    [DANGER] Danger Zone: Permanently delete all uploaded documents 
    and clear their embeddings from the vector store.
    """
    try:
        # Import vector store
        from app.services.vectorstore import get_vector_store, reset_vector_store
        
        vector_store = get_vector_store()
        
        # Count before deleting
        total_docs = db.query(Document).count()
        total_chunks = db.query(DocumentChunk).count()
        vector_count = vector_store.get_count()
        
        logger.info(f"[STATS] Before deletion: {total_docs} docs, {total_chunks} chunks, {vector_count} vectors")
        
        # 🧠 Clear vector store FIRST and reset singleton
        try:
            vector_store.clear()
            
            # CRITICAL: Reset the singleton instance to create a fresh vector store
            reset_vector_store()
            
            # Get the new instance and verify it's empty
            vector_store = get_vector_store()
            after_clear = vector_store.get_count()
            
            logger.info(f"[SUCCESS] Vector store cleared and reset: {vector_count} -> {after_clear}")
            
            if after_clear > 0:
                raise Exception(f"Vector store not fully cleared: {after_clear} items remain")
            
        except Exception as e:
            logger.error(f"[ERROR] Vector store clear failed: {str(e)}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"Failed to clear vector store: {str(e)}")
        
        # 🧹 Delete from database (this cascades to chunks)
        deleted_docs = db.query(Document).delete()
        db.commit()
        logger.info(f"[SUCCESS] Database cleared: {deleted_docs} documents deleted")
        
        # Final verification
        remaining_docs = db.query(Document).count()
        remaining_chunks = db.query(DocumentChunk).count()
        remaining_vectors = vector_store.get_count()
        
        logger.info(f"[STATS] After deletion: {remaining_docs} docs, {remaining_chunks} chunks, {remaining_vectors} vectors")
        
        if remaining_docs > 0 or remaining_chunks > 0 or remaining_vectors > 0:
            logger.warning(f"[WARN] Incomplete deletion detected!")
            return {
                "status": "partial_success",
                "message": "Some data may not have been fully cleared",
                "documents_deleted": total_docs,
                "chunks_deleted": total_chunks,
                "vectors_deleted": vector_count,
                "remaining": {
                    "documents": remaining_docs,
                    "chunks": remaining_chunks,
                    "vectors": remaining_vectors
                }
            }
        
        return {
            "status": "success",
            "message": f"Successfully cleared all data: {total_docs} documents, {total_chunks} chunks, {vector_count} vectors",
            "documents_deleted": total_docs,
            "chunks_deleted": total_chunks,
            "vectors_deleted": vector_count,
            "verification": {
                "remaining_documents": 0,
                "remaining_chunks": 0,
                "remaining_vectors": 0
            }
        }
    
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"[ERROR] Failed to clear documents: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to clear documents: {str(e)}")

# ============================================================
# 📄 Document Management
# ============================================================
@router.get("/documents")
async def list_documents(skip: int = 0, limit: int = 100, db: Session = Depends(get_db)):
    try:
        documents = db.query(Document).offset(skip).limit(limit).all()
        total_count = db.query(Document).count()
        return {
            "total": total_count,
            "count": len(documents),
            "documents": [
                {
                    "id": d.id,
                    "filename": d.filename,
                    "content_type": d.content_type,
                    "size": d.size,
                    "status": d.status,
                    "chunks_count": d.chunks_count,
                    "uploaded_at": d.uploaded_at.isoformat() if d.uploaded_at else None,
                    "processed_at": d.processed_at.isoformat() if d.processed_at else None
                }
                for d in documents
            ]
        }
    except Exception as e:
        logger.error(f"[ERROR] Failed to fetch documents: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/documents/{document_id}")
async def get_document(document_id: str, db: Session = Depends(get_db)):
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return {
        "id": document.id,
        "filename": document.filename,
        "content_type": document.content_type,
        "size": document.size,
        "status": document.status,
        "chunks_count": document.chunks_count,
        "uploaded_at": document.uploaded_at,
        "processed_at": document.processed_at,
        "metadata": document.metadata
    }


@router.delete("/documents/{document_id}")
async def delete_document(document_id: str, db: Session = Depends(get_db)):
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    filename = document.filename
    db.delete(document)
    db.commit()
    return {"status": "success", "message": f"Document '{filename}' deleted successfully"}


# ============================================================
# 📝 Query History
# ============================================================
@router.get("/queries")
async def list_queries(limit: int = 10, skip: int = 0, db: Session = Depends(get_db)):
    queries = db.query(QueryModel).order_by(QueryModel.created_at.desc()).offset(skip).limit(limit).all()
    total = db.query(QueryModel).count()
    return {
        "total": total,
        "count": len(queries),
        "queries": [
            {
                "id": q.id,
                "query": q.query_text,
                "answer": (q.answer[:200] + "...") if q.answer and len(q.answer) > 200 else q.answer,
                "strategy": q.strategy_used,
                "processing_time": q.processing_time,
                "confidence_score": q.confidence_score,
                "created_at": q.created_at.isoformat() if q.created_at else None
            }
            for q in queries
        ]
    }


@router.get("/queries/{query_id}")
async def get_query(query_id: str, db: Session = Depends(get_db)):
    query = db.query(QueryModel).filter(QueryModel.id == query_id).first()
    if not query:
        raise HTTPException(status_code=404, detail="Query not found")
    return {
        "id": query.id,
        "query_text": query.query_text,
        "answer": query.answer,
        "strategy_used": query.strategy_used,
        "processing_time": query.processing_time,
        "confidence_score": query.confidence_score,
        "created_at": query.created_at
    }


# ============================================================
# 🧪 Health & Stats
# ============================================================
@router.get("/health", response_model=HealthCheckResponse)
async def health_check(db: Session = Depends(get_db)):
    """
    Comprehensive health check that tests all critical services.
    """

    from sqlalchemy import text

    # ✅ Initialize all components with default "unknown"
    components = {
        "database": "unknown",
        "llm_service": "unknown",
        "embedding_service": "unknown",
        "vector_store": "unknown",
        "knowledge_graph": "unknown",
        "memory": "unknown",
        "agents": "unknown",
    }

    all_healthy = True

    # 1. 🧪 Test Database
    try:
        db.execute(text("SELECT 1"))
        components["database"] = "operational"
    except Exception as e:
        components["database"] = f"error: {str(e)}"
        all_healthy = False

    # 2. 🤖 Test LLM Service (Groq)
    rag_service = None
    try:
        rag_service = get_rag_service()
        test_response = await rag_service.llm_service.generate(
            "Say 'OK' if you can read this.",
            temperature=0.1,
            max_tokens=10
        )

        if test_response and len(test_response) > 0:
            components["llm_service"] = "operational"
        else:
            components["llm_service"] = "error: empty response"
            all_healthy = False

    except Exception as e:
        components["llm_service"] = f"error: {str(e)}"
        all_healthy = False

    # 3. 🧠 Test Embedding Service
    try:
        if rag_service:
            test_embedding = rag_service.embedding_service.embed_text("test")
            if test_embedding and len(test_embedding) > 0:
                components["embedding_service"] = "operational"
            else:
                components["embedding_service"] = "error: no embedding generated"
                all_healthy = False
    except Exception as e:
        components["embedding_service"] = f"error: {str(e)}"
        all_healthy = False

    # 4. 🧰 Test Vector Store
    try:
        if rag_service:
            _ = rag_service.vectorstore.get_count()
            components["vector_store"] = "operational"
    except Exception as e:
        components["vector_store"] = f"error: {str(e)}"
        all_healthy = False

    # 5. 🕸 Test Knowledge Graph
    try:
        if rag_service and hasattr(rag_service, "graph_builder") and rag_service.graph_builder:
            components["knowledge_graph"] = "operational"
        else:
            components["knowledge_graph"] = "disabled"
    except Exception as e:
        components["knowledge_graph"] = f"error: {str(e)}"

    # 6. 🧠 Memory & Agents (Basic Check)
    try:
        components["memory"] = "operational"
        components["agents"] = "operational"
    except Exception as e:
        components["memory"] = f"error: {str(e)}"
        components["agents"] = f"error: {str(e)}"

    # 📊 Overall status
    status = "healthy" if all_healthy else "degraded"

    return HealthCheckResponse(
        status=status,
        timestamp=datetime.now(),
        version="2.0.0",
        components=components
    )


@router.get("/health/detailed")
async def detailed_health_check(db: Session = Depends(get_db)):
    """
    Detailed health check with performance metrics
    """
    from sqlalchemy import text
    import time
    
    results = {
        "overall_status": "checking...",
        "timestamp": datetime.now().isoformat(),
        "checks": {}
    }
    
    # Database Check
    db_start = time.time()
    try:
        db.execute(text("SELECT 1"))
        results["checks"]["database"] = {
            "status": "healthy",
            "response_time_ms": round((time.time() - db_start) * 1000, 2),
            "message": "Database connection successful"
        }
    except Exception as e:
        results["checks"]["database"] = {
            "status": "unhealthy",
            "response_time_ms": round((time.time() - db_start) * 1000, 2),
            "error": str(e)
        }
    
    # LLM Service Check (Groq)
    llm_start = time.time()
    try:
        rag_service = get_rag_service()
        test_response = await rag_service.llm_service.generate(
            "Respond with exactly: 'Health check OK'",
            temperature=0,
            max_tokens=20
        )
        
        results["checks"]["llm_service"] = {
            "status": "healthy",
            "response_time_ms": round((time.time() - llm_start) * 1000, 2),
            "model": os.getenv("LLM_MODEL", "llama-3.1-8b-instant"),
            "provider": "Groq",
            "test_response": test_response[:50] if test_response else None,
            "message": "LLM API responding correctly"
        }
    except Exception as e:
        results["checks"]["llm_service"] = {
            "status": "unhealthy",
            "response_time_ms": round((time.time() - llm_start) * 1000, 2),
            "error": str(e),
            "message": "Failed to connect to Groq API"
        }
    
    # Embedding Service Check
    embed_start = time.time()
    try:
        test_text = "Health check test"
        embedding = rag_service.embedding_service.embed_text(test_text)
        
        results["checks"]["embedding_service"] = {
            "status": "healthy",
            "response_time_ms": round((time.time() - embed_start) * 1000, 2),
            "model": os.getenv("EMBEDDING_MODEL", "sentence-transformers/all-MiniLM-L6-v2"),
            "embedding_dimension": len(embedding) if embedding else 0,
            "message": "Embedding service working"
        }
    except Exception as e:
        results["checks"]["embedding_service"] = {
            "status": "unhealthy",
            "response_time_ms": round((time.time() - embed_start) * 1000, 2),
            "error": str(e)
        }
    
    # Vector Store Check
    vector_start = time.time()
    try:
        count = rag_service.vectorstore.get_count()
        
        # Try a test search
        test_embedding = rag_service.embedding_service.embed_text("test query")
        search_results = rag_service.vectorstore.search(test_embedding, top_k=1)
        
        results["checks"]["vector_store"] = {
            "status": "healthy",
            "response_time_ms": round((time.time() - vector_start) * 1000, 2),
            "total_documents": count,
            "search_working": len(search_results) >= 0,
            "message": f"Vector store operational with {count} documents"
        }
    except Exception as e:
        results["checks"]["vector_store"] = {
            "status": "unhealthy",
            "response_time_ms": round((time.time() - vector_start) * 1000, 2),
            "error": str(e)
        }
    
    # RAG Pipeline End-to-End Test
    rag_start = time.time()
    try:
        # Only test if we have documents
        if rag_service.vectorstore.get_count() > 0:
            test_result = await rag_service.execute_query(
                query="test health check",
                top_k=1,
                strategy=RAGStrategy.SIMPLE
            )
            
            results["checks"]["rag_pipeline"] = {
                "status": "healthy",
                "response_time_ms": round((time.time() - rag_start) * 1000, 2),
                "answer_generated": len(test_result.get("answer", "")) > 0,
                "chunks_retrieved": len(test_result.get("retrieved_chunks", [])),
                "message": "Full RAG pipeline working"
            }
        else:
            results["checks"]["rag_pipeline"] = {
                "status": "ready",
                "response_time_ms": 0,
                "message": "No documents indexed yet - upload documents to test"
            }
    except Exception as e:
        results["checks"]["rag_pipeline"] = {
            "status": "unhealthy",
            "response_time_ms": round((time.time() - rag_start) * 1000, 2),
            "error": str(e)
        }
    
    # Determine overall status
    unhealthy_count = sum(
        1 for check in results["checks"].values() 
        if check.get("status") == "unhealthy"
    )
    
    if unhealthy_count == 0:
        results["overall_status"] = "healthy"
    elif unhealthy_count <= 2:
        results["overall_status"] = "degraded"
    else:
        results["overall_status"] = "unhealthy"
    
    # Add system stats
    results["statistics"] = {
        "total_documents": db.query(Document).count(),
        "total_queries": db.query(QueryModel).count(),
        "total_chunks": rag_service.vectorstore.get_count(),
        "active_sessions": db.query(SessionModel).filter(SessionModel.is_active == True).count()
    }
    
    return results


@router.get("/health/quick")
async def quick_health_check():
    """
    Quick health check without testing external services
    Useful for load balancers
    """
    return {
        "status": "ok",
        "timestamp": datetime.now().isoformat(),
        "service": "RAG System",
        "version": "2.0.0"
    }


@router.get("/stats")
async def get_statistics(db: Session = Depends(get_db)):
    from sqlalchemy import func
    total_docs = db.query(Document).count()
    total_queries = db.query(QueryModel).count()
    total_chunks = db.query(func.sum(Document.chunks_count)).scalar() or 0
    avg_time = db.query(func.avg(QueryModel.processing_time)).scalar() or 0
    strategy_stats = db.query(
        QueryModel.strategy_used,
        func.count(QueryModel.id)
    ).group_by(QueryModel.strategy_used).all()

    return {
        "total_documents": total_docs,
        "total_queries": total_queries,
        "total_chunks": int(total_chunks),
        "average_processing_time": float(avg_time),
        "strategy_distribution": {s: c for s, c in strategy_stats}
    }


# ============================================================
# 💬 Session Endpoints
# ============================================================
@router.post("/sessions")
async def create_session(user_id: str = None, db: Session = Depends(get_db)):
    session = SessionModel(
        id=str(uuid.uuid4()),
        user_id=user_id,
        started_at=datetime.utcnow(),
        last_activity=datetime.utcnow(),
        is_active=True
    )
    db.add(session)
    db.commit()
    db.refresh(session)
    return {
        "session_id": session.id,
        "user_id": session.user_id,
        "started_at": session.started_at.isoformat(),
        "status": "active"
    }


@router.get("/sessions/{session_id}")
async def get_session(session_id: str, db: Session = Depends(get_db)):
    session = db.query(SessionModel).filter(SessionModel.id == session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    queries = db.query(QueryModel).filter(
        QueryModel.session_id == session_id
    ).order_by(QueryModel.created_at.asc()).all()

    return {
        "session_id": session.id,
        "user_id": session.user_id,
        "started_at": session.started_at.isoformat(),
        "last_activity": session.last_activity.isoformat(),
        "is_active": session.is_active,
        "queries": [
            {
                "id": q.id,
                "query": q.query_text,
                "answer": q.answer,
                "timestamp": q.created_at.isoformat() if q.created_at else None
            }
            for q in queries
        ]
    }


# ============================================================
# 🧪 Test
# ============================================================
@router.get("/test")
async def test_endpoint():
    return {"message": "[SUCCESS] RAG router is working!"}


@router.get("/debug/vectorstore")
async def debug_vectorstore():
    """[DEBUG] Debug endpoint to check vector store state"""
    try:
        from app.services.vectorstore import get_vector_store
        
        vector_store = get_vector_store()
        
        return {
            "status": "ok",
            "vector_count": vector_store.get_count(),
            "faiss_index_total": vector_store.index.ntotal,
            "dimension": vector_store.dimension,
            "documents_sample": vector_store.documents[:3] if len(vector_store.documents) > 0 else [],
            "metadata_sample": vector_store.metadata[:3] if len(vector_store.metadata) > 0 else [],
            "documents_list_length": len(vector_store.documents),
            "metadata_list_length": len(vector_store.metadata)
        }
    except Exception as e:
        logger.error(f"[ERROR] Debug endpoint failed: {str(e)}", exc_info=True)
        return {
            "status": "error",
            "error": str(e),
            "traceback": traceback.format_exc()
        }
    
    
# In routers/rag_router.py

@router.get("/graph/{document_id}")
async def get_document_graph(document_id: str, db: Session = Depends(get_db)):
    """
    Get graph structure for a document
    """
    entities = db.query(GraphEntity).filter(
        GraphEntity.properties["document_id"].astext == document_id
    ).all()
    
    relationships = db.query(GraphRelationship).join(
        GraphEntity, GraphEntity.id == GraphRelationship.source_id
    ).filter(
        GraphEntity.properties["document_id"].astext == document_id
    ).all()
    
    return {
        "nodes": [
            {
                "id": e.id,
                "name": e.name,
                "type": e.type,
                "description": e.description
            } for e in entities
        ],
        "edges": [
            {
                "source": r.source_id,
                "target": r.target_id,
                "relation": r.relation_type,
                "weight": r.weight
            } for r in relationships
        ]
    }